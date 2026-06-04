#!/usr/bin/env python3
"""Replace company-level meeting content and requirement sections in a DOCX template."""

from __future__ import annotations

import argparse
import copy
import re
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


SECTION_RE = re.compile(r"^#{1,2}\s*(会议内容|会议要求)\s*$")
SUBHEADING_RE = re.compile(r"^#{2,3}\s*(.+?)\s*$")
CHINESE_NUMERALS = "一二三四五六七八九十"


def normalize_heading(text: str) -> str:
    return text.strip().lstrip("一二三四五六七八九十、.． ").strip()


def find_section_start(doc: Document, normalized_name: str) -> int:
    for index, paragraph in enumerate(doc.paragraphs):
        if normalize_heading(paragraph.text) == normalized_name:
            return index
    raise ValueError(f"Section heading not found: {normalized_name}")


def find_end_after(doc: Document, start: int, end_marker: str) -> int:
    for index in range(start + 1, len(doc.paragraphs)):
        if doc.paragraphs[index].text.strip().startswith(end_marker):
            return index
    raise ValueError(f"End marker not found after section: {end_marker}")


def remove_paragraphs(doc: Document, start_exclusive: int, end_exclusive: int):
    for paragraph in list(doc.paragraphs[start_exclusive + 1 : end_exclusive]):
        paragraph._element.getparent().remove(paragraph._element)


def clear_runs(paragraph_element):
    for child in list(paragraph_element):
        if child.tag == qn("w:r"):
            paragraph_element.remove(child)


def make_run(text: str, template_run=None, bold: bool | None = None):
    run = OxmlElement("w:r")
    if template_run is not None and template_run._r.rPr is not None:
        run.append(copy.deepcopy(template_run._r.rPr))
    if bold is not None:
        rpr = run.find(qn("w:rPr"))
        if rpr is None:
            rpr = OxmlElement("w:rPr")
            run.insert(0, rpr)
        existing = rpr.find(qn("w:b"))
        if bold and existing is None:
            rpr.append(OxmlElement("w:b"))
        if not bold and existing is not None:
            rpr.remove(existing)
    text_element = OxmlElement("w:t")
    if text.startswith(" ") or text.endswith(" "):
        text_element.set(qn("xml:space"), "preserve")
    text_element.text = text
    run.append(text_element)
    return run


def clone_paragraph_with_text(template_paragraph, text: str, bold: bool | None = None):
    element = copy.deepcopy(template_paragraph._p)
    clear_runs(element)
    template_runs = [run for run in template_paragraph.runs if run.text]
    first_run = template_runs[0] if template_runs else None
    element.append(make_run(text, first_run, bold=bold))
    return element


def split_content(path: Path) -> dict[str, list[tuple[str, str]]]:
    current_section = None
    sections: dict[str, list[tuple[str, str]]] = {"会议内容": [], "会议要求": []}

    for raw_line in path.read_text(encoding="utf-8").splitlines():
        line = raw_line.strip()
        if not line:
            continue

        section_match = SECTION_RE.match(line)
        if section_match:
            current_section = section_match.group(1)
            continue

        if current_section is None:
            raise ValueError("Content must start with '# 会议内容' or '# 会议要求'.")

        subheading_match = SUBHEADING_RE.match(line)
        if subheading_match:
            sections[current_section].append(("heading", subheading_match.group(1)))
        else:
            sections[current_section].append(("body", line))

    missing = [name for name, items in sections.items() if not items]
    if missing:
        raise ValueError(f"Missing replacement content for: {', '.join(missing)}")

    return sections


def insert_items_before(end_paragraph, items: list[tuple[str, str]], body_template, heading_template):
    for kind, text in items:
        if kind == "heading":
            element = clone_paragraph_with_text(heading_template, text, bold=True)
        else:
            element = clone_paragraph_with_text(body_template, text)
        end_paragraph._p.addprevious(element)


def number_requirement_headings(items: list[tuple[str, str]]) -> list[tuple[str, str]]:
    heading_index = 0
    numbered_items = []
    for kind, text in items:
        if kind == "heading":
            heading_index += 1
            if not text.startswith("（"):
                if heading_index <= len(CHINESE_NUMERALS):
                    numeral = CHINESE_NUMERALS[heading_index - 1]
                else:
                    numeral = str(heading_index)
                text = f"（{numeral}） {text}"
        numbered_items.append((kind, text))
    return numbered_items


def replace_company_sections(template: Path, content: Path, output: Path):
    doc = Document(str(template))
    sections = split_content(content)

    content_start = find_section_start(doc, "会议内容")
    requirements_start = find_section_start(doc, "会议要求")
    supervision_start = find_end_after(doc, requirements_start, "督办")

    content_template = doc.paragraphs[content_start + 1]
    requirements_body_template = doc.paragraphs[requirements_start + 1]
    requirements_heading_template = next(
        (
            paragraph
            for paragraph in doc.paragraphs[requirements_start + 1 : supervision_start]
            if any(run.bold for run in paragraph.runs)
        ),
        requirements_body_template,
    )

    supervision_paragraph = doc.paragraphs[supervision_start]
    remove_paragraphs(doc, requirements_start, supervision_start)
    insert_items_before(
        supervision_paragraph,
        number_requirement_headings(sections["会议要求"]),
        requirements_body_template,
        requirements_heading_template,
    )

    requirements_start = find_section_start(doc, "会议要求")
    requirements_paragraph = doc.paragraphs[requirements_start]
    remove_paragraphs(doc, content_start, requirements_start)
    insert_items_before(
        requirements_paragraph,
        sections["会议内容"],
        content_template,
        content_template,
    )

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output))


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--template", type=Path, required=True)
    parser.add_argument("--content", type=Path, required=True)
    parser.add_argument("--output", type=Path, required=True)
    args = parser.parse_args()

    replace_company_sections(args.template, args.content, args.output)


if __name__ == "__main__":
    main()
