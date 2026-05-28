#!/usr/bin/env python3
"""Replace the body of a meeting brief section while preserving DOCX formatting."""

from __future__ import annotations

import argparse
import copy
import re
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


NUMBERED_LABEL_RE = re.compile(r"^(\d+[.．、]\s*[^：:]{1,40}[：:])(.+)$")


def iter_nonempty_paragraphs(doc: Document):
    for index, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip():
            yield index, paragraph


def find_section_bounds(doc: Document, section_heading: str, end_marker: str):
    section_index = None
    for index, paragraph in iter_nonempty_paragraphs(doc):
        if paragraph.text.strip() == section_heading:
            section_index = index
            break
    if section_index is None:
        raise ValueError(f"Section heading not found: {section_heading}")

    end_index = None
    for index in range(section_index + 1, len(doc.paragraphs)):
        if doc.paragraphs[index].text.strip().startswith(end_marker):
            end_index = index
            break
    if end_index is None:
        raise ValueError(f"End marker not found after section: {end_marker}")
    if end_index <= section_index + 1:
        raise ValueError("Section has no replaceable body paragraphs.")

    return section_index, end_index


def paragraph_signature(paragraph):
    runs = [run for run in paragraph.runs if run.text]
    first_run = runs[0] if runs else None
    return {
        "bold": first_run.bold if first_run is not None else None,
        "font": first_run.font.name if first_run is not None else None,
        "text": paragraph.text.strip(),
    }


def choose_style_paragraphs(doc: Document, start: int, end: int):
    subsection = None
    body = None
    bold_body = None

    for paragraph in doc.paragraphs[start + 1 : end]:
        text = paragraph.text.strip()
        if not text:
            continue
        sig = paragraph_signature(paragraph)
        if subsection is None and sig["bold"] and "楷体" in str(sig["font"]):
            subsection = paragraph
        if body is None and "仿宋" in str(sig["font"]):
            body = paragraph
        if bold_body is None and "仿宋" in str(sig["font"]) and any(run.bold for run in paragraph.runs):
            bold_body = paragraph

    if subsection is None:
        subsection = doc.paragraphs[start + 1]
    if body is None:
        body = doc.paragraphs[min(start + 2, end - 1)]
    if bold_body is None:
        bold_body = body

    return subsection, body, bold_body


def remove_paragraphs(doc: Document, start_exclusive: int, end_exclusive: int):
    for paragraph in list(doc.paragraphs[start_exclusive + 1 : end_exclusive]):
        paragraph._element.getparent().remove(paragraph._element)


def clear_runs(paragraph_element):
    for child in list(paragraph_element):
        if child.tag == qn("w:r"):
            paragraph_element.remove(child)


def make_run(text: str, template_run=None, bold=None):
    run = OxmlElement("w:r")
    if template_run is not None and template_run._r.rPr is not None:
        run.append(copy.deepcopy(template_run._r.rPr))
    if bold is not None:
        rpr = run.find(qn("w:rPr"))
        if rpr is None:
            rpr = OxmlElement("w:rPr")
            run.insert(0, rpr)
        existing = rpr.find(qn("w:b"))
        if bold and existing is None:
            rpr.append(OxmlElement("w:b"))
        if not bold and existing is not None:
            rpr.remove(existing)
    t = OxmlElement("w:t")
    if text.startswith(" ") or text.endswith(" "):
        t.set(qn("xml:space"), "preserve")
    t.text = text
    run.append(t)
    return run


def clone_paragraph_with_text(template_paragraph, text: str, kind: str):
    element = copy.deepcopy(template_paragraph._p)
    clear_runs(element)

    template_runs = [run for run in template_paragraph.runs if run.text]
    first_run = template_runs[0] if template_runs else None
    normal_run = next((run for run in template_runs if not run.bold), first_run)
    bold_run = next((run for run in template_runs if run.bold), first_run)

    if kind == "body":
        match = NUMBERED_LABEL_RE.match(text)
        if match:
            label, rest = match.groups()
            element.append(make_run(label, bold_run, bold=True))
            element.append(make_run(rest, normal_run, bold=False))
        else:
            element.append(make_run(text, normal_run))
    else:
        element.append(make_run(text, first_run, bold=True))

    return element


def parse_content(path: Path):
    items = []
    for raw_line in path.read_text(encoding="utf-8").splitlines():
        line = raw_line.strip()
        if not line:
            continue
        if line.startswith("##"):
            title = line.lstrip("#").strip()
            if title:
                items.append(("heading", title))
        else:
            items.append(("body", line))
    if not items:
        raise ValueError("Replacement content is empty.")
    return items


def replace_section(template: Path, content: Path, output: Path, section_heading: str, end_marker: str):
    doc = Document(str(template))
    start, end = find_section_bounds(doc, section_heading, end_marker)
    subsection_template, body_template, bold_body_template = choose_style_paragraphs(doc, start, end)
    end_paragraph = doc.paragraphs[end]

    items = parse_content(content)
    remove_paragraphs(doc, start, end)

    for kind, text in items:
        source = subsection_template if kind == "heading" else bold_body_template if NUMBERED_LABEL_RE.match(text) else body_template
        new_element = clone_paragraph_with_text(source, text, "heading" if kind == "heading" else "body")
        end_paragraph._p.addprevious(new_element)

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output))


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--template", type=Path, required=True)
    parser.add_argument("--content", type=Path, required=True)
    parser.add_argument("--output", type=Path, required=True)
    parser.add_argument("--section-heading", default="会议重点讨论事项")
    parser.add_argument("--end-marker", default="承办部门：")
    args = parser.parse_args()

    replace_section(args.template, args.content, args.output, args.section_heading, args.end_marker)


if __name__ == "__main__":
    main()
