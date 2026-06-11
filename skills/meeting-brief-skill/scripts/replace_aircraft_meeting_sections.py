#!/usr/bin/env python3
"""Replace 飞机端 monthly brief sections while preserving DOCX formatting."""

from __future__ import annotations

import argparse
import copy
import re
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


SECTION_RE = re.compile(r"^#\s*(会议重点内容|参会领导作工作指示)\s*$")
HEADING_RE = re.compile(r"^(#{2,3})\s*(.+?)\s*$")
NUMBERED_LABEL_RE = re.compile(r"^(\d+[.．、]\s*[^：:]{1,40}[：:])(.+)$")
CHINESE_NUMERALS = "一二三四五六七八九十"


def normalize_heading(text: str) -> str:
    return text.strip().lstrip("一二三四五六七八九十、.． ").strip()


def paragraph_text(paragraph) -> str:
    return paragraph.text.strip()


def find_section_start(doc: Document, normalized_name: str) -> int:
    for index, paragraph in enumerate(doc.paragraphs):
        if normalize_heading(paragraph_text(paragraph)) == normalized_name:
            return index
    raise ValueError(f"Section heading not found: {normalized_name}")


def find_end_after(doc: Document, start: int, normalized_or_prefix: str) -> int:
    for index in range(start + 1, len(doc.paragraphs)):
        text = paragraph_text(doc.paragraphs[index])
        if not text:
            continue
        if normalize_heading(text) == normalized_or_prefix or text.startswith(normalized_or_prefix):
            return index
    raise ValueError(f"End marker not found after section: {normalized_or_prefix}")


def remove_paragraphs(doc: Document, start_exclusive: int, end_exclusive: int):
    for paragraph in list(doc.paragraphs[start_exclusive + 1 : end_exclusive]):
        paragraph._element.getparent().remove(paragraph._element)


def clear_runs(paragraph_element):
    for child in list(paragraph_element):
        if child.tag == qn("w:r"):
            paragraph_element.remove(child)


def run_fonts(run) -> str:
    if run._r.rPr is None:
        return ""
    fonts = run._r.rPr.find(qn("w:rFonts"))
    if fonts is None:
        return ""
    return " ".join(fonts.attrib.values())


def paragraph_fonts(paragraph) -> str:
    return " ".join(run_fonts(run) for run in paragraph.runs)


def has_bold(paragraph) -> bool:
    return any(run.bold for run in paragraph.runs)


def pick_template(paragraphs, *, font_contains: str, bold: bool | None = None):
    for paragraph in paragraphs:
        text = paragraph_text(paragraph)
        if not text:
            continue
        if font_contains in paragraph_fonts(paragraph):
            if bold is None or has_bold(paragraph) == bold or (bold and has_bold(paragraph)):
                return paragraph
    return None


def choose_templates(doc: Document, start: int, end: int):
    section_paragraphs = doc.paragraphs[start + 1 : end]
    topic = pick_template(section_paragraphs, font_contains="楷体", bold=True)
    body = pick_template(section_paragraphs, font_contains="仿宋", bold=False)
    bold_body = pick_template(section_paragraphs, font_contains="仿宋", bold=True)

    if topic is None:
        topic = doc.paragraphs[start + 1]
    if body is None:
        body = doc.paragraphs[min(start + 1, end - 1)]
    if bold_body is None:
        bold_body = body
    return topic, body, bold_body


def make_run(text: str, template_run=None, bold: bool | None = None):
    run = OxmlElement("w:r")
    if template_run is not None and template_run._r.rPr is not None:
        run.append(copy.deepcopy(template_run._r.rPr))
    if bold is not None:
        rpr = run.find(qn("w:rPr"))
        if rpr is None:
            rpr = OxmlElement("w:rPr")
            run.insert(0, rpr)
        existing = rpr.find(qn("w:b"))
        if bold and existing is None:
            rpr.append(OxmlElement("w:b"))
        if not bold and existing is not None:
            rpr.remove(existing)
    text_element = OxmlElement("w:t")
    if text.startswith(" ") or text.endswith(" "):
        text_element.set(qn("xml:space"), "preserve")
    text_element.text = text
    run.append(text_element)
    return run


def clone_paragraph_with_text(template_paragraph, text: str, kind: str):
    element = copy.deepcopy(template_paragraph._p)
    clear_runs(element)

    template_runs = [run for run in template_paragraph.runs if run.text]
    first_run = template_runs[0] if template_runs else None
    normal_run = next((run for run in template_runs if not run.bold), first_run)
    bold_run = next((run for run in template_runs if run.bold), first_run)

    if kind == "body":
        match = NUMBERED_LABEL_RE.match(text)
        if match:
            label, rest = match.groups()
            element.append(make_run(label, bold_run, bold=True))
            element.append(make_run(rest, normal_run, bold=False))
        else:
            element.append(make_run(text, normal_run))
    else:
        element.append(make_run(text, first_run, bold=True))
    return element


def split_content(path: Path) -> dict[str, list[tuple[str, str]]]:
    sections: dict[str, list[tuple[str, str]]] = {
        "会议重点内容": [],
        "参会领导作工作指示": [],
    }
    current_section = None

    for raw_line in path.read_text(encoding="utf-8").splitlines():
        line = raw_line.strip()
        if not line:
            continue

        section_match = SECTION_RE.match(line)
        if section_match:
            current_section = section_match.group(1)
            continue
        if current_section is None:
            raise ValueError("Content must start with '# 会议重点内容' or '# 参会领导作工作指示'.")

        heading_match = HEADING_RE.match(line)
        if heading_match:
            level, title = heading_match.groups()
            kind = "topic" if level == "##" else "subheading"
            sections[current_section].append((kind, title))
        else:
            sections[current_section].append(("body", line))

    missing = [name for name, items in sections.items() if not items]
    if missing:
        raise ValueError(f"Missing replacement content for: {', '.join(missing)}")
    return sections


def number_topics(items: list[tuple[str, str]]) -> list[tuple[str, str]]:
    topic_index = 0
    numbered_items = []
    for kind, text in items:
        if kind == "topic":
            topic_index += 1
            if not text.startswith("（"):
                numeral = CHINESE_NUMERALS[topic_index - 1] if topic_index <= len(CHINESE_NUMERALS) else str(topic_index)
                text = f"（{numeral}）{text}"
        numbered_items.append((kind, text))
    return numbered_items


def insert_items_before(end_paragraph, items, topic_template, body_template, bold_body_template):
    for kind, text in items:
        if kind in {"topic", "subheading"}:
            source = topic_template
            clone_kind = "heading"
        else:
            source = bold_body_template if NUMBERED_LABEL_RE.match(text) else body_template
            clone_kind = "body"
        end_paragraph._p.addprevious(clone_paragraph_with_text(source, text, clone_kind))


def replace_aircraft_sections(template: Path, content: Path, output: Path):
    doc = Document(str(template))
    sections = split_content(content)

    focus_start = find_section_start(doc, "会议重点内容")
    instruction_start = find_section_start(doc, "参会领导作工作指示")
    footer_start = find_end_after(doc, instruction_start, "承办部门：")

    focus_topic, focus_body, focus_bold_body = choose_templates(doc, focus_start, instruction_start)
    instruction_topic, instruction_body, instruction_bold_body = choose_templates(doc, instruction_start, footer_start)

    footer_paragraph = doc.paragraphs[footer_start]
    remove_paragraphs(doc, instruction_start, footer_start)
    insert_items_before(
        footer_paragraph,
        sections["参会领导作工作指示"],
        instruction_topic,
        instruction_body,
        instruction_bold_body,
    )

    instruction_start = find_section_start(doc, "参会领导作工作指示")
    instruction_paragraph = doc.paragraphs[instruction_start]
    remove_paragraphs(doc, focus_start, instruction_start)
    insert_items_before(
        instruction_paragraph,
        number_topics(sections["会议重点内容"]),
        focus_topic,
        focus_body,
        focus_bold_body,
    )

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output))


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--template", type=Path, required=True)
    parser.add_argument("--content", type=Path, required=True)
    parser.add_argument("--output", type=Path, required=True)
    args = parser.parse_args()

    replace_aircraft_sections(args.template, args.content, args.output)


if __name__ == "__main__":
    main()
